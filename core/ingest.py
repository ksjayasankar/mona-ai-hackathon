"""Turn any file into Claude content blocks.

Gemini reads PDFs and images natively (vision), so we deliberately avoid a heavy OCR
stack. DOCX/XLSX get a light text extraction. Returns a list of content blocks ready
to pass to core.llm.ask/extract.
"""
from __future__ import annotations

import base64
from pathlib import Path

IMAGE_TYPES = {".png": "image/png", ".jpg": "image/jpeg", ".jpeg": "image/jpeg", ".webp": "image/webp", ".gif": "image/gif"}


def _b64(data: bytes) -> str:
    return base64.standard_b64encode(data).decode("ascii")


def file_to_blocks(path: str | Path) -> list[dict]:
    """One file -> content blocks. Handles pdf, png/jpg, docx, xlsx, txt/eml."""
    path = Path(path)
    data = path.read_bytes()
    return bytes_to_blocks(data, path.suffix.lower(), path.name)


def bytes_to_blocks(data: bytes, suffix: str, name: str = "upload") -> list[dict]:
    suffix = suffix.lower()
    if suffix == ".pdf":
        return [{"type": "document", "source": {"type": "base64", "media_type": "application/pdf", "data": _b64(data)}}]
    if suffix in IMAGE_TYPES:
        return [{"type": "image", "source": {"type": "base64", "media_type": IMAGE_TYPES[suffix], "data": _b64(data)}}]
    if suffix == ".docx":
        return [{"type": "text", "text": f"[{name} — extracted text]\n{_docx_text(data)}"}]
    if suffix in (".xlsx", ".xlsm"):
        return [{"type": "text", "text": f"[{name} — spreadsheet]\n{_xlsx_text(data)}"}]
    # txt, eml, md, csv, anything else -> best-effort text
    return [{"type": "text", "text": f"[{name}]\n{data.decode('utf-8', errors='replace')}"}]


def _docx_text(data: bytes) -> str:
    import io
    from docx import Document

    doc = Document(io.BytesIO(data))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(c.text for c in row.cells))
    return "\n".join(parts)


def _xlsx_text(data: bytes, max_rows: int = 200) -> str:
    import io
    from openpyxl import load_workbook

    wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    out = []
    for ws in wb.worksheets:
        out.append(f"## Sheet: {ws.title}")
        for i, row in enumerate(ws.iter_rows(values_only=True)):
            if i >= max_rows:
                out.append("... (truncated)")
                break
            out.append(" | ".join("" if v is None else str(v) for v in row))
    return "\n".join(out)
